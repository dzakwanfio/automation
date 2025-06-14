import datetime
import json
import logging
import os
import re
import subprocess

from django.conf import settings
import jwt
import openpyxl
import pandas as pd
from django.conf import settings
from django.contrib import messages
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.core.mail import send_mail
from django.http import FileResponse, Http404, HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse
from django.utils import timezone
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST
from docx import Document
from io import BytesIO
import traceback
from copy import deepcopy
import time
import uuid

# Setup logging untuk debugging
logging.basicConfig(
    filename='process_files.log',
    level=logging.DEBUG,
    format="%(asctime)s - %(levelname)s - %(message)s",
)

from .forms import OtomatisasiForm
from .models import LogHistory, Otomatisasi, UploadedFile, Siswa

@login_required(login_url="login")
def homepage(request):
    return render(request, "homepage.html")

def user_login(request):
    if request.method == "POST":
        email = request.POST.get("email")
        password = request.POST.get("password")

        user = authenticate(request, username=email, password=password)
        if user is not None:
            login(request, user)
            return redirect("homepage")

        messages.error(request, "Email atau password salah")

    return render(request, "login.html")

def register(request):
    if request.method == "POST":
        email = request.POST.get("email")
        password = request.POST.get("password")
        confirm_password = request.POST.get("confirm_password")

        if password != confirm_password:
            messages.error(request, "Password dan konfirmasi password tidak cocok")
            return redirect("register")

        if len(password) < 8:
            messages.error(request, "Password harus setidaknya 8 karakter")
            return redirect("register")

        if User.objects.filter(email=email).exists():
            messages.error(request, "Email sudah terdaftar")
            return redirect("register")

        user = User.objects.create_user(username=email, email=email, password=password)
        user.is_active = False
        user.save()

        token = generate_verification_token(email)

        verification_link = request.build_absolute_uri(
            reverse("verify_email", kwargs={"token": token})
        )
        email_subject = "Verifikasi Akun Anda"
        email_body = f"Klik link berikut untuk mengaktifkan akun Anda: {verification_link}"

        send_mail(email_subject, email_body, "your_email@gmail.com", [email], fail_silently=False)

        messages.success(request, "Email verifikasi telah dikirim. Silakan cek inbox Anda.")
        return redirect("login")

    return render(request, "register.html")

def generate_verification_token(email):
    payload = {
        "email": email,
        "exp": datetime.datetime.utcnow() + datetime.timedelta(hours=24),
        "iat": datetime.datetime.utcnow(),
    }
    return jwt.encode(payload, settings.SECRET_KEY, algorithm="HS256")

def verify_email(request, token):
    try:
        payload = jwt.decode(token, settings.SECRET_KEY, algorithms=["HS256"])
        email = payload["email"]

        user = get_object_or_404(User, email=email)

        if user.is_active:
            messages.error(request, "Email sudah diverifikasi sebelumnya.")
        else:
            user.is_active = True
            user.save()
            messages.success(request, "Email berhasil diverifikasi! Silakan login.")

        return redirect("login")

    except jwt.ExpiredSignatureError:
        messages.error(request, "Token telah kedaluwarsa. Silakan daftar kembali.")
    except jwt.DecodeError:
        messages.error(request, "Token tidak valid.")

    return redirect("register")

logger = logging.getLogger(__name__)

@login_required(login_url="login")
def input_data(request):
    errors = []

    if request.method == "POST":
        course_name = request.POST.get("course_name")
        start_date = request.POST.get("start_date")
        end_date = request.POST.get("end_date")
        course_model = request.POST.get("course_model")
        destination = request.POST.get("destination")
        upload_file = request.FILES.get("file_upload")

        required_columns = [
            "No", "Nama", "Jenis_Kelamin", "NIK", "Tempat_Lahir", "Tanggal_Lahir", "NISN",
            "Agama_LKP", "Handphone", "Kewarganegaraan", "Jenis_Tinggal", "Tanggal_Masuk",
            "Email", "Nama_Ortu", "NIK_Ortu", "Pekerjaan_Ortu", "Pendidikan_Ortu",
            "Penghasilan_Ortu", "Handphone_Ortu", "Tempat_Lahir_Ortu", "Tanggal_Lahir_Ortu",
            "Asal", "Alamat", "RT", "RW", "Kecamatan", "Kelurahan", "Kab/Kota", "Propinsi",
            "Nama_Ibu_kandung", "Nama_Ayah", "Agama_Kemdikbud", "Penerima_KPS", "Layak_PIP",
            "Penerima_KIP", "Kode_Pos", "Jenis_tinggal", "Alat_Transportasi",
            "Pendidikan_Terakhir", "Nama_Lembaga", "Jabatan", "Alamat_Kantor", "Telp_Kantor"
        ]

        if upload_file:
            try:
                # Baca file Excel tanpa memaksa dtype=str
                df = pd.read_excel(upload_file, sheet_name="all")
                
                # Log tipe data yang dibaca
                logger.debug("[DEBUG] Tipe data setiap kolom: %s", df.dtypes.to_dict())

                # Bersihkan spasi di kolom string
                for column in df.columns:
                    if df[column].dtype == 'object':  # Hanya untuk kolom string
                        df[column] = df[column].astype(str).str.strip()

                # Periksa kolom yang hilang
                missing_columns = [col for col in required_columns if col not in df.columns]
                if missing_columns:
                    errors.append(f"Kolom berikut tidak ditemukan: {', '.join(missing_columns)}")
                    return render(request, "input_data.html", {"errors": errors})

                # Log nilai mentah untuk kolom tanggal sebelum parsing
                date_columns = ["Tanggal_Lahir", "Tanggal_Masuk", "Tanggal_Lahir_Ortu"]
                for col in date_columns:
                    if col in df.columns:
                        logger.debug("[DEBUG] Nilai mentah %s: %s", col, df[col].tolist())

                # Konversi kolom tanggal
                for col in date_columns:
                    if col in df.columns:
                        # Jika kolom sudah bertipe datetime, ekstrak tanggal langsung
                        if pd.api.types.is_datetime64_any_dtype(df[col]):
                            df[col] = df[col].dt.date
                        else:
                            # Bersihkan nilai yang tidak valid
                            df[col] = df[col].astype(str).str.strip()
                            df[col] = df[col].replace(["N/A", "", "nan"], pd.NA)
                            # Coba parsing tanggal
                            df[col] = pd.to_datetime(df[col], dayfirst=True, errors='coerce')
                            # Ekstrak hanya tanggal
                            df[col] = df[col].apply(lambda x: x.date() if pd.notnull(x) else None)

                        # Log nilai setelah parsing
                        logger.debug("[DEBUG] Nilai setelah parsing %s: %s", col, df[col].tolist())

                        # Validasi format
                        invalid_dates = df[col][df[col].apply(lambda x: x is not None and not isinstance(x, datetime.date))]
                        if not invalid_dates.empty:
                            errors.append(f'Nilai "{invalid_dates.iloc[0]}" mempunyai format tanggal yang tidak valid. Tanggal harus dalam format DD/MM/YYYY.')
                            return render(request, "input_data.html", {"errors": errors})

                # Periksa baris kosong
                df_required = df[required_columns].replace(["", " ", "nan"], pd.NA)
                empty_rows = df_required[df_required.isna().any(axis=1)]
                if not empty_rows.empty:
                    empty_row_indices = empty_rows.index + 2
                    logger.debug("[DEBUG] Baris kosong ditemukan di indeks: %s", empty_row_indices.tolist())
                    for index in empty_row_indices:
                        empty_columns = df_required.iloc[index - 2][df_required.iloc[index - 2].isna()].index.tolist()
                        logger.debug("[DEBUG] Kolom kosong di baris %d: %s", index, empty_columns)
                        errors.append(f"Baris {index} pada file memiliki sel kosong pada kolom: {', '.join(empty_columns)}")
                    return render(request, "input_data.html", {"errors": errors})

                # Simpan ke UploadedFile
                uploaded_file = UploadedFile.objects.create(
                    course_name=course_name,
                    start_date=start_date,
                    end_date=end_date,
                    course_model=course_model,
                    destination=destination,
                    file=upload_file,
                )

                # Simpan data Excel ke model Peserta
                for _, row in df.iterrows():
                    Peserta.objects.create(
                        uploaded_file=uploaded_file,
                        nama=row["Nama"] if pd.notna(row["Nama"]) else None,
                        jenis_kelamin=row["Jenis_Kelamin"] if pd.notna(row["Jenis_Kelamin"]) else None,
                        nik=row["NIK"] if pd.notna(row["NIK"]) else None,
                        tempat_lahir=row["Tempat_Lahir"] if pd.notna(row["Tempat_Lahir"]) else None,
                        tanggal_lahir=row["Tanggal_Lahir"] if pd.notna(row["Tanggal_Lahir"]) else None,
                        nisn=row["NISN"] if pd.notna(row["NISN"]) else None,
                        agama_lkp=row["Agama_LKP"] if pd.notna(row["Agama_LKP"]) else None,
                        handphone=row["Handphone"] if pd.notna(row["Handphone"]) else None,
                        kewarganegaraan=row["Kewarganegaraan"] if pd.notna(row["Kewarganegaraan"]) else None,
                        jenis_tinggal=row["Jenis_Tinggal"] if pd.notna(row["Jenis_Tinggal"]) else None,
                        tanggal_masuk=row["Tanggal_Masuk"] if pd.notna(row["Tanggal_Masuk"]) else None,
                        email=row["Email"] if pd.notna(row["Email"]) else None,
                        nama_ortu=row["Nama_Ortu"] if pd.notna(row["Nama_Ortu"]) else None,
                        nik_ortu=row["NIK_Ortu"] if pd.notna(row["NIK_Ortu"]) else None,
                        pekerjaan_ortu=row["Pekerjaan_Ortu"] if pd.notna(row["Pekerjaan_Ortu"]) else None,
                        pendidikan_ortu=row["Pendidikan_Ortu"] if pd.notna(row["Pendidikan_Ortu"]) else None,
                        penghasilan_ortu=row["Penghasilan_Ortu"] if pd.notna(row["Penghasilan_Ortu"]) else None,
                        handphone_ortu=row["Handphone_Ortu"] if pd.notna(row["Handphone_Ortu"]) else None,
                        tempat_lahir_ortu=row["Tempat_Lahir_Ortu"] if pd.notna(row["Tempat_Lahir_Ortu"]) else None,
                        tanggal_lahir_ortu=row["Tanggal_Lahir_Ortu"] if pd.notna(row["Tanggal_Lahir_Ortu"]) else None,
                        asal=row["Asal"] if pd.notna(row["Asal"]) else None,
                        alamat=row["Alamat"] if pd.notna(row["Alamat"]) else None,
                        rt=row["RT"] if pd.notna(row["RT"]) else None,
                        rw=row["RW"] if pd.notna(row["RW"]) else None,
                        kecamatan=row["Kecamatan"] if pd.notna(row["Kecamatan"]) else None,
                        kelurahan=row["Kelurahan"] if pd.notna(row["Kelurahan"]) else None,
                        kab_kota=row["Kab/Kota"] if pd.notna(row["Kab/Kota"]) else None,
                        propinsi=row["Propinsi"] if pd.notna(row["Propinsi"]) else None,
                        nama_ibu_kandung=row["Nama_Ibu_kandung"] if pd.notna(row["Nama_Ibu_kandung"]) else None,
                        nama_ayah=row["Nama_Ayah"] if pd.notna(row["Nama_Ayah"]) else None,
                        agama_kemdikbud=row["Agama_Kemdikbud"] if pd.notna(row["Agama_Kemdikbud"]) else None,
                        penerima_kps=row["Penerima_KPS"] if pd.notna(row["Penerima_KPS"]) else None,
                        layak_pip=row["Layak_PIP"] if pd.notna(row["Layak_PIP"]) else None,
                        penerima_kip=row["Penerima_KIP"] if pd.notna(row["Penerima_KIP"]) else None,
                        kode_pos=row["Kode_Pos"] if pd.notna(row["Kode_Pos"]) else None,
                        alat_transportasi=row["Alat_Transportasi"] if pd.notna(row["Alat_Transportasi"]) else None,
                        pendidikan_terakhir=row["Pendidikan_Terakhir"] if pd.notna(row["Pendidikan_Terakhir"]) else None,
                        nama_lembaga=row["Nama_Lembaga"] if pd.notna(row["Nama_Lembaga"]) else None,
                        jabatan=row["Jabatan"] if pd.notna(row["Jabatan"]) else None,
                        alamat_kantor=row["Alamat_Kantor"] if pd.notna(row["Alamat_Kantor"]) else None,
                        telp_kantor=row["Telp_Kantor"] if pd.notna(row["Telp_Kantor"]) else None,
                        kota=row["Kab/Kota"] if pd.notna(row["Kab/Kota"]) else None,
                        is_converted=False
                    )

                messages.success(request, "File berhasil diunggah dan data peserta disimpan!")
                return redirect("input_data")

            except Exception as e:
                errors.append(f"File tidak sesuai. Error: {str(e)}")
                logger.error(f"Error processing file: {str(e)}")
                return render(request, "input_data.html", {"errors": errors})

        else:
            errors.append("Silakan upload file.")
            return render(request, "input_data.html", {"errors": errors})

    logger.info(f"Loading input_data page with GET request. URL: {request.path}, Referer: {request.META.get('HTTP_REFERER', 'Unknown')}")
    return render(request, "input_data.html", {"errors": errors})

@login_required(login_url="login")
def upload_page(request):
    return render(request, "upload.html")

@login_required(login_url="login")
def otomatisasi(request):
    files = UploadedFile.objects.all()  # Ambil semua file yang belum dihapus
    logging.info(f"Files in otomatisasi view: {files.count()} files found")  # Tambah logging
    return render(request, "otomatisasi.html", {"files": files})

from django.utils.timezone import localtime

@login_required(login_url="login")
def log_history(request):
    logs = LogHistory.objects.all()
    for log in logs:
        log.local_time = localtime(log.upload_date).strftime("%d %b %Y %H:%M")
    return render(request, "log_history.html", {"logs": logs})

@login_required(login_url="login")
def logoutview(request):
    logout(request)
    return redirect("login")

def forgot_pw(request):
    return render(request, "forgot_pw.html")

def forgot_password_notification(request):
    return render(request, "forgot_pwnotif.html")

def reset_password(request):
    if request.method == "POST":
        password = request.POST.get("password")
        confirm_password = request.POST.get("confirm_password")

        if password == confirm_password:
            messages.success(request, "Password berhasil direset! Silakan login.")
            return redirect("login")
        else:
            messages.error(request, "Password tidak cocok. Silakan coba lagi.")

    return render(request, "reset_password.html")

@login_required(login_url="login")
def delete_otomatisasi(request, id):  # Ubah file_id menjadi id
    if request.method == "POST":
        try:
            file = UploadedFile.objects.get(id=id)
            file_path = file.file.path
            file.delete()
            if os.path.exists(file_path):
                os.remove(file_path)
                logging.info(f"File deleted from disk: {file_path}")
            logging.info(f"Deleted UploadedFile entry with ID {id}")
            return JsonResponse({"status": "success", "message": "File deleted successfully"})
        except UploadedFile.DoesNotExist:
            logging.info(f"UploadedFile with ID {id} already deleted or not found")
            return JsonResponse({"status": "error", "message": "File not found"}, status=404)
        except Exception as e:
            logging.error(f"Error deleting file with ID {id}: {str(e)}")
            return JsonResponse({"status": "error", "message": f"Error deleting file: {str(e)}"}, status=500)
    else:
        logging.warning("Invalid method in delete_otomatisasi")
        return JsonResponse({"status": "error", "message": "Metode tidak diizinkan, gunakan POST"}, status=405)

@login_required(login_url="login")
def edit_otomatisasi(request, id):
    file_obj = get_object_or_404(UploadedFile, pk=id)
    if request.method == 'POST':
        form = OtomatisasiForm(request.POST, request.FILES, instance=file_obj)
        if form.is_valid():
            form.save()
            messages.success(request, "Data berhasil diperbarui!")
            return redirect('otomatisasi')
    else:
        form = OtomatisasiForm(instance=file_obj)
    return render(request, 'edit_otomatisasi.html', {'form': form, 'file_obj': file_obj})

@login_required(login_url="login")
def upload_data(request):
    if request.method == "POST":
        course_name = request.POST.get("course_name")
        course_model = request.POST.get("course_model")

        if course_name and course_model:
            UploadedFile.objects.create(course_name=course_name, course_model=course_model)
            messages.success(request, "Data berhasil ditambahkan!")
            return redirect("otomatisasi")

    return render(request, "upload.html")

@login_required(login_url="login")
def download_file(request, file_id):
    file_obj = get_object_or_404(UploadedFile, id=file_id)

    file_path = file_obj.file.path  
    if not os.path.exists(file_path):
        raise Http404("File tidak ditemukan.")

    response = FileResponse(open(file_path, "rb"))
    response["Content-Disposition"] = f'attachment; filename="{os.path.basename(file_path)}"'
    return response

import json
import logging
import os
import shutil
import subprocess

from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse
from django.utils import timezone
from django.views.decorators.csrf import csrf_exempt

from .models import LogHistory, UploadedFile, Siswa

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s', handlers=[logging.FileHandler('process_files.log')])

@csrf_exempt
@login_required(login_url="login")
def process_files(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            file_ids = data.get("file_ids", [])

            logging.info(f"process_files called with file_ids: {file_ids}, from URL: {request.path}, Referer: {request.META.get('HTTP_REFERER', 'Unknown')}")

            if not file_ids:
                logging.warning("No files selected in process_files")
                return JsonResponse({"status": "error", "message": "No files selected.", "last_row": 0}, status=400)

            files = UploadedFile.objects.filter(id__in=file_ids)
            if not files.exists():
                logging.warning("No valid files found for the given IDs")
                return JsonResponse({"status": "error", "message": "No valid files found.", "last_row": 0}, status=400)

            file_path_dict = {file.id: file.file.path for file in files}
            script_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'otomatisasi.py')
            
            if not os.path.exists(script_path):
                logging.error(f"Script not found at: {script_path}")
                return JsonResponse({
                    "status": "error",
                    "message": "Script otomatisasi tidak ditemukan",
                    "last_row": 0,
                    "detail": f"Path yang dicari: {script_path}"
                }, status=400)

            processed_files = []
            failed_files = []
            last_row = 0
            status = "success"
            message = f"{len(files)} file berhasil diproses!"

            temp_dir = os.path.join(settings.MEDIA_ROOT, 'temp_failed_files')
            if not os.path.exists(temp_dir):
                os.makedirs(temp_dir)

            for file in files:
                file_path = file.file.path
                destination_url = file.destination  # Ambil URL dari field destination
                logging.info(f"Starting process for file: {file_path}, exists: {os.path.exists(file_path)}, URL: {destination_url}")
                
                # Validasi destination_url
                if not destination_url:
                    logging.error(f"No destination URL provided for file: {file_path}")
                    LogHistory.objects.create(
                        name=os.path.basename(file.file.name),
                        upload_date=timezone.now(),
                        course_name=file.course_name,
                        status='Failed (No destination URL)',
                        process_time=timezone.now(),
                        file_path=os.path.basename(file.file.name),
                        file_id=file.id
                    )
                    file.is_failed = True
                    file.save()
                    file.delete()
                    failed_files.append(file.id)
                    message = f"Process failed for file {os.path.basename(file_path)}: No destination URL provided"
                    status = "error"
                    continue

                process = None
                try:
                    # Tambahkan --url ke perintah
                    process = subprocess.Popen(
                        ["python", script_path, file_path, "--url", destination_url],
                        stdout=subprocess.PIPE,
                        stderr=subprocess.PIPE,
                        text=True,
                        cwd=os.path.dirname(os.path.abspath(__file__))
                    )

                    stdout, stderr = process.communicate(timeout=600)

                    logging.debug(f"Subprocess stdout for {file_path}: {stdout}")
                    logging.debug(f"Subprocess stderr for {file_path}: {stderr}")
                    logging.debug(f"Subprocess returncode for {file_path}: {process.returncode}")

                    script_output = {}
                    if stdout.strip():
                        try:
                            script_output = json.loads(stdout.strip())
                        except json.JSONDecodeError as e:
                            logging.error(f"JSON decode error for {file_path}: {str(e)}")
                            script_output = {"status": "error", "message": "Gagal memproses output dari script otomatisasi"}

                    file_status = script_output.get("status", "error")
                    file_message = script_output.get("message", "Terjadi kesalahan saat memroses file")
                    last_row = script_output.get("last_row", 0)

                    temp_file_path = file_path
                    if file_status != "success":
                        file_name = os.path.basename(file_path)
                        temp_file_path = os.path.join(temp_dir, file_name)
                        shutil.copy2(file_path, temp_file_path)
                        logging.info(f"File copied to temp location: {temp_file_path}")

                    LogHistory.objects.create(
                        name=os.path.basename(file.file.name),
                        upload_date=timezone.now(),
                        course_name=file.course_name,
                        status='Success' if file_status == "success" else f'Failed (Stopped at row {last_row})',
                        process_time=timezone.now(),
                        file_path=os.path.basename(file.file.name),
                        file_id=file.id
                    )

                    if file_status == "success":
                        try:
                            if os.path.exists(file_path):
                                os.remove(file_path)
                                logging.info(f"File deleted after successful processing: {file_path}")
                        except Exception as e:
                            logging.warning(f"Failed to delete file {file_path}: {e}")
                        file.delete()
                        logging.info(f"Removed UploadedFile entry for successful file: {file_path}")
                        processed_files.append(file.id)
                    else:
                        file.is_failed = True
                        file.last_processed_row = last_row
                        file.save()
                        file.delete()
                        logging.info(f"Removed UploadedFile entry for failed file: {file_path}, file preserved at: {temp_file_path}")
                        failed_files.append(file.id)
                        message = f"Process failed at file {os.path.basename(file_path)}: {file_message} (Stopped at row {last_row})"
                        status = "error"
                        break

                except subprocess.TimeoutExpired as e:
                    logging.error(f"Timeout processing file {file_path}: {str(e)}")
                    if process:
                        process.kill()
                        logging.info(f"Process for {file_path} has been killed due to timeout.")
                    file_name = os.path.basename(file_path)
                    temp_file_path = os.path.join(temp_dir, file_name)
                    shutil.copy2(file_path, temp_file_path)
                    logging.info(f"File copied to temp location: {temp_file_path}")

                    LogHistory.objects.create(
                        name=os.path.basename(file.file.name),
                        upload_date=timezone.now(),
                        course_name=file.course_name,
                        status='Failed (Timeout)',
                        process_time=timezone.now(),
                        file_path=os.path.basename(file.file.name),
                        file_id=file.id
                    )
                    file.is_failed = True
                    file.last_processed_row = last_row
                    file.save()
                    file.delete()
                    logging.info(f"Removed UploadedFile entry for timed-out file: {file_path}, file preserved at: {temp_file_path}")
                    failed_files.append(file.id)
                    message = f"Timeout processing file {os.path.basename(file_path)}: Process took too long"
                    status = "error"
                    break
                except Exception as e:
                    logging.error(f"Unexpected error processing {file_path}: {str(e)}")
                    if process:
                        process.kill()
                        logging.info(f"Process for {file_path} has been killed due to error.")
                    file_name = os.path.basename(file_path)
                    temp_file_path = os.path.join(temp_dir, file_name)
                    shutil.copy2(file_path, temp_file_path)
                    logging.info(f"File copied to temp location: {temp_file_path}")

                    LogHistory.objects.create(
                        name=os.path.basename(file.file.name),
                        upload_date=timezone.now(),
                        course_name=file.course_name,
                        status='Failed (Unexpected error)',
                        process_time=timezone.now(),
                        file_path=os.path.basename(file.file.name),
                        file_id=file.id
                    )
                    file.is_failed = True
                    file.last_processed_row = last_row
                    file.save()
                    file.delete()
                    logging.info(f"Removed UploadedFile entry for failed file with unexpected error: {file_path}, file preserved at: {temp_file_path}")
                    failed_files.append(file.id)
                    message = f"Unexpected error processing {os.path.basename(file_path)}: {str(e)}"
                    status = "error"
                    break
                finally:
                    if process:
                        try:
                            process.kill()
                            logging.info(f"Ensured process for {file_path} is terminated.")
                        except:
                            pass

            remaining_files = [f for f in files if f.id not in processed_files and f.id not in failed_files]
            for remaining_file in remaining_files:
                logging.info(f"Retained unprocessed file: {remaining_file.file.path}")

            logging.info(f"Returning response: status={status}, failed_file_ids={failed_files}")
            return JsonResponse({
                "status": status,
                "message": message,
                "last_row": last_row,
                "failed_file_ids": failed_files if failed_files else []
            }, status=200)

        except json.JSONDecodeError as e:
            logging.error(f"JSON decode error in request body: {str(e)}")
            return JsonResponse({"status": "error", "message": "Data request tidak valid", "last_row": 0}, status=400)
        except Exception as e:
            logging.error(f"Unexpected error in process_files: {str(e)}")
            return JsonResponse({"status": "error", "message": f"Error server: {str(e)}", "last_row": 0}, status=500)

    logging.warning("Invalid method in process_files")
    return JsonResponse({"status": "error", "message": "Metode tidak diizinkan", "last_row": 0}, status=405)

@csrf_exempt
@login_required(login_url="login")
def resume_process(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            log_id = data.get("log_id")

            logging.info(f"resume_process called with log_id: {log_id}, from URL: {request.path}, Referer: {request.META.get('HTTP_REFERER', 'Unknown')}")

            if not log_id:
                logging.warning("No log_id provided in resume_process")
                return JsonResponse({"status": "error", "message": "No log ID provided.", "last_row": 0})

            log_entry = get_object_or_404(LogHistory, id=log_id)
            match = re.search(r"Stopped at row (\d+)", log_entry.status)
            if not match:
                logging.error(f"Could not extract last row from status: {log_entry.status}")
                return JsonResponse({
                    "status": "error",
                    "message": "Cannot determine the last processed row.",
                    "last_row": 0
                })

            last_row = int(match.group(1))
            if last_row < 1:
                last_row = 2

            # Ambil file terkait dari log_entry
            file_path = os.path.join(settings.MEDIA_ROOT, log_entry.file_path)
            logging.info(f"Resuming with file path: {file_path}, exists: {os.path.exists(file_path)}")
            if not os.path.exists(file_path):
                temp_file_path = os.path.join(settings.MEDIA_ROOT, 'temp_failed_files', log_entry.file_path)
                if os.path.exists(temp_file_path):
                    file_path = temp_file_path
                    logging.info(f"File found in temp location: {file_path}")
                else:
                    logging.error(f"File not found at: {file_path} or {temp_file_path}")
                    return JsonResponse({
                        "status": "error",
                        "message": f"File tidak ditemukan di server.",
                        "last_row": 0
                    })

            # Ambil destination dari UploadedFile terkait
            try:
                uploaded_file = UploadedFile.objects.get(id=log_entry.file_id)
                destination_url = uploaded_file.destination
                logging.info(f"Destination URL for resume: {destination_url}")
            except UploadedFile.DoesNotExist:
                logging.error(f"No UploadedFile found for log_id: {log_id}")
                return JsonResponse({
                    "status": "error",
                    "message": "Tidak dapat menemukan data file terkait untuk melanjutkan proses.",
                    "last_row": 0
                })

            if not destination_url:
                logging.error(f"No destination URL provided for log_id: {log_id}")
                return JsonResponse({
                    "status": "error",
                    "message": "Tidak ada URL destinasi yang tersedia untuk melanjutkan proses.",
                    "last_row": 0
                })

            script_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'otomatisasi.py')
            if not os.path.exists(script_path):
                logging.error(f"Script not found at: {script_path}")
                return JsonResponse({
                    "status": "error",
                    "message": "Script otomatisasi tidak ditemukan",
                    "last_row": 0,
                    "detail": f"Path yang dicari: {script_path}"
                })

            process = None
            try:
                process = subprocess.Popen(
                    ["python", script_path, file_path, "--resume-from", str(last_row), "--url", destination_url],
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    text=True,
                    cwd=os.path.dirname(os.path.abspath(__file__))
                )

                stdout, stderr = process.communicate(timeout=600)

                logging.debug(f"Resume subprocess stdout: {stdout}")
                logging.debug(f"Resume subprocess stderr: {stderr}")
                logging.debug(f"Resume subprocess returncode: {process.returncode}")

                script_output = {}
                if stdout.strip():
                    try:
                        script_output = json.loads(stdout.strip())
                    except json.JSONDecodeError as e:
                        logging.error(f"JSON decode error in resume_process: {str(e)}")
                        return JsonResponse({
                            "status": "error",
                            "message": "Gagal memproses output dari script otomatisasi",
                            "last_row": last_row,
                            "detail": stderr if stderr else "No stderr output"
                        })

                status = script_output.get("status", "error")
                message = script_output.get("message", "Terjadi kesalahan saat melanjutkan proses")
                last_row = script_output.get("last_row", last_row)

                if status == "success":
                    log_entry.status = "Success"
                    log_entry.process_time = timezone.now()
                    log_entry.save()

                    if os.path.exists(file_path):
                        os.remove(file_path)
                        logging.info(f"File deleted from disk after successful resume: {file_path}")

                    logging.info(f"Successfully resumed processing for log ID {log_id}")
                    return JsonResponse({
                        "status": "success",
                        "message": "Proses berhasil dilanjutkan dan selesai!",
                        "last_row": last_row
                    })
                else:
                    log_entry.status = f"Failed (Stopped at row {last_row})"
                    log_entry.process_time = timezone.now()
                    log_entry.save()

                    logging.error(f"Resume process failed at row {last_row}: {message}")
                    return JsonResponse({
                        "status": "error",
                        "message": message,
                        "last_row": last_row,
                        "detail": stderr if stderr else "No stderr output"
                    })

            except subprocess.TimeoutExpired as e:
                logging.error(f"Timeout in resume_process for log_id {log_id}: {str(e)}")
                if process:
                    process.kill()
                    logging.info(f"Process for log_id {log_id} has been killed due to timeout.")
                log_entry.status = "Failed (Timeout)"
                log_entry.process_time = timezone.now()
                log_entry.save()
                return JsonResponse({
                    "status": "error",
                    "message": "Timeout: Resume process took too long",
                    "last_row": last_row
                })
            except Exception as e:
                logging.error(f"Unexpected error in resume_process: {str(e)}")
                if process:
                    process.kill()
                    logging.info(f"Process for log_id {log_id} has been killed due to error.")
                return JsonResponse({"status": "error", "message": str(e), "last_row": 0})
            finally:
                if process:
                    try:
                        process.kill()
                        logging.info(f"Ensured process for log_id {log_id} is terminated.")
                    except:
                        pass

        except Exception as e:
            logging.error(f"Unexpected error in resume_process: {str(e)}")
            return JsonResponse({"status": "error", "message": str(e), "last_row": 0})

    logging.warning("Invalid method in resume_process")
    return JsonResponse({"status": "error", "message": "Metode tidak diizinkan", "last_row": 0})

@login_required(login_url="login")
def data_siswa(request):
    siswa_list = Siswa.objects.all()  # Ambil semua data siswa dari model
    return render(request, 'data_siswa.html', {'siswa_list': siswa_list})

@require_POST
@login_required(login_url="login")
def delete_data_siswa_single(request, id):
    try:
        siswa = Siswa.objects.get(id=id)
        siswa.delete()
        logging.info(f"Deleted Siswa entry with ID {id}")
        return JsonResponse({"status": "success", "message": "Data siswa berhasil dihapus!"})
    except Siswa.DoesNotExist:
        logging.info(f"Siswa with ID {id} not found")
        return JsonResponse({"status": "error", "message": "Data siswa tidak ditemukan"}, status=404)
    except Exception as e:
        logging.error(f"Error deleting Siswa with ID {id}: {str(e)}")
        return JsonResponse({"status": "error", "message": f"Error deleting data: {str(e)}"}, status=500)

@require_POST
@login_required(login_url="login")
def delete_data_siswa(request):
    try:
        data = json.loads(request.body)
        siswa_ids = data.get('siswa_ids', [])
        if not siswa_ids:
            return JsonResponse({"status": "error", "message": "Tidak ada data siswa yang dipilih."}, status=400)
        
        deleted_count = Siswa.objects.filter(id__in=siswa_ids).delete()[0]
        logging.info(f"Deleted {deleted_count} Siswa entries with IDs {siswa_ids}")
        return JsonResponse({"status": "success", "message": f"{deleted_count} data siswa berhasil dihapus!"})
    except json.JSONDecodeError:
        return JsonResponse({"status": "error", "message": "Data request tidak valid"}, status=400)
    except Exception as e:
        logging.error(f"Error deleting multiple Siswa entries: {str(e)}")
        return JsonResponse({"status": "error", "message": f"Error deleting data: {str(e)}"}, status=500)

@login_required(login_url="login")
def edit_data_siswa(request, id):
    siswa = get_object_or_404(Siswa, pk=id)
    if request.method == 'POST':
        # Asumsi Anda memiliki form untuk Siswa (buat form jika belum ada)
        siswa.nama = request.POST.get('nama', siswa.nama)
        siswa.nikp = request.POST.get('nikp', siswa.nikp)
        siswa.jenis_kelamin = request.POST.get('jenis_kelamin', siswa.jenis_kelamin)
        siswa.alamat = request.POST.get('alamat', siswa.alamat)
        siswa.nomor_hp = request.POST.get('nomor_hp', siswa.nomor_hp)
        siswa.save()
        messages.success(request, "Data siswa berhasil diperbarui!")
        return redirect('data_siswa')
    return render(request, 'edit_data_siswa.html', {'siswa': siswa})

from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from .models import Peserta
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


@login_required(login_url="login")
def generate_document(request):
    peserta_list = Peserta.objects.all()
    logger.info(f"Mengambil {peserta_list.count()} data peserta untuk generate_document")
    return render(request, 'generate_document.html', {'files': peserta_list})

@login_required(login_url="login")
def delete_peserta(request):
    logger.info("[DEBUG] Memproses permintaan delete_peserta pada: %s", request.path)
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            peserta_ids = data.get("peserta_ids", [])
            logger.debug("[DEBUG] ID peserta yang diterima: %s", peserta_ids)

            if not peserta_ids:
                logger.warning("[WARN] Tidak ada peserta yang dipilih untuk dihapus")
                return JsonResponse({"status": "error", "message": "Tidak ada peserta yang dipilih."}, status=400)

            # Konversi ID ke integer
            peserta_ids = [int(id) for id in peserta_ids]
            logger.debug("[DEBUG] ID peserta setelah konversi: %s", peserta_ids)

            # Verifikasi ID
            existing_peserta = Peserta.objects.filter(id__in=peserta_ids)
            if not existing_peserta.exists():
                logger.warning("[WARN] Tidak ada peserta ditemukan dengan ID: %s", peserta_ids)
                return JsonResponse({"status": "error", "message": "Data peserta tidak ditemukan."}, status=404)

            # Hapus peserta
            deleted_count, _ = existing_peserta.delete()
            logger.info("[INFO] Berhasil menghapus %s peserta dengan ID: %s", deleted_count, peserta_ids)

            return JsonResponse({"status": "success", "message": f"{deleted_count} data peserta berhasil dihapus!"})
        except json.JSONDecodeError as e:
            logger.error("[ERROR] Error parsing JSON: %s", str(e))
            return JsonResponse({"status": "error", "message": "Data permintaan tidak valid."}, status=400)
        except ValueError as e:
            logger.error("[ERROR] Error konversi ID: %s", str(e))
            return JsonResponse({"status": "error", "message": "ID peserta tidak valid."}, status=400)
        except Exception as e:
            logger.error("[ERROR] Error menghapus peserta: %s", str(e))
            return JsonResponse({"status": "error", "message": f"Error: {str(e)}"}, status=500)
    else:
        logger.warning("[WARN] Metode tidak diizinkan: %s", request.method)
        return JsonResponse({"status": "error", "message": "Metode tidak diizinkan."}, status=405)

@login_required(login_url="login")
def convert_document(request):
    logger.info("[DEBUG] Memproses permintaan convert_document pada: %s", request.path)
    if request.method == "POST":
        try:
            # Langkah 1: Parse request body
            logger.debug("[DEBUG] Request body: %s", request.body)
            data = json.loads(request.body)
            peserta_ids = data.get("peserta_ids", [])
            jadwal = data.get("jadwal")
            tuk = data.get("tuk")
            skema = data.get("skema")
            asesor = data.get("asesor")
            lokasi_sertif = data.get("lokasi_sertif")
            logger.debug("[DEBUG] Data diterima: peserta_ids=%s, jadwal=%s, tuk=%s, skema=%s, asesor=%s, lokasi_sertif=%s",
                         peserta_ids, jadwal, tuk, skema, asesor, lokasi_sertif)

            # Langkah 2: Validasi input
            if not peserta_ids:
                logger.warning("[WARN] Tidak ada peserta dipilih untuk konversi")
                return JsonResponse({"status": "error", "message": "Pilih setidaknya satu peserta."}, status=400)

            if not all([jadwal, tuk, skema, asesor, lokasi_sertif]):
                logger.warning("[WARN] Field form tidak lengkap")
                return JsonResponse({"status": "error", "message": "Semua field form harus diisi."}, status=400)

            # Langkah 3: Ambil data peserta
            logger.debug("[DEBUG] Mengambil peserta dengan ID: %s", peserta_ids)
            peserta_list = Peserta.objects.filter(id__in=peserta_ids)
            if not peserta_list.exists():
                logger.warning("[WARN] Peserta tidak ditemukan: %s", peserta_ids)
                return JsonResponse({"status": "error", "message": "Peserta tidak ditemukan."}, status=404)

            # Langkah 4: Format Tanggal_Sertif
            tanggal_sertif = format_tanggal_indonesia(timezone.now())
            logger.debug("[DEBUG] Tanggal_Sertif: %s", tanggal_sertif)

            # Langkah 5: Pemetaan skema ke folder template
            skema_to_folder = {
                "Associate Data Analyst": "folder1",
                "Instruktur Junior (KKNI Level III)": "folder2",
                "Junior Information Management": "folder3",
                "Pemasangan Jaringan Komputer": "folder4",
                "Pengelolaan Backup dan Restore Data": "folder5",
                "Pengelolaan Data Aplikasi Perkantoran": "folder6",
                "Pengelolaan Keamanan Data Pengguna": "folder7",
                "Pengelolaan Keamanan Jaringan": "folder8",
            }

            if skema not in skema_to_folder:
                logger.error("[ERROR] Skema tidak valid: %s", skema)
                return JsonResponse({"status": "error", "message": f"Skema '{skema}' tidak valid."}, status=400)

            folder_name = skema_to_folder[skema]
            logger.info("[INFO] Menggunakan template dari folder: %s untuk skema: %s", folder_name, skema)

            # Langkah 6: Path ke template Word berdasarkan folder
            template_paths = [
                os.path.join(os.path.dirname(__file__), 'templates', 'docx', folder_name, 'FormPendaftaran.docx'),
                os.path.join(os.path.dirname(__file__), 'templates', 'docx', folder_name, 'AK01.docx'),
                os.path.join(os.path.dirname(__file__), 'templates', 'docx', folder_name, 'APL02.docx'),
                os.path.join(os.path.dirname(__file__), 'templates', 'docx', folder_name, 'APL01.docx')
            ]
            for path in template_paths:
                if not os.path.exists(path):
                    logger.error("[ERROR] Template Word tidak ditemukan di: %s", path)
                    return JsonResponse({"status": "error", "message": f"Template Word {path} tidak ditemukan."}, status=500)

            # Buat direktori untuk menyimpan file sementara di STATICFILES_DIRS
            temp_dir = os.path.join(settings.STATICFILES_DIRS[0], 'temp')
            if not os.path.exists(temp_dir):
                os.makedirs(temp_dir)
                logger.info("[INFO] Membuat direktori temp: %s", temp_dir)

            # Bersihkan file lama
            now = time.time()
            for filename in os.listdir(temp_dir):
                file_path = os.path.join(temp_dir, filename)
                if os.path.exists(file_path) and os.stat(file_path).st_mtime < now - 3600:  # 1 jam
                    os.remove(file_path)
                    logger.info("[INFO] Menghapus file lama: %s", file_path)

            # Langkah 7: Proses semua peserta dan buat dokumen untuk kedua template
            download_urls = []
            temp_files = []
            counter = 1  # Inisialisasi counter untuk nomor urut

            for template_path in template_paths:
                final_doc = Document()
                first_doc = True

                for index, peserta in enumerate(peserta_list):
                    logger.debug("[DEBUG] Memproses peserta: %s (ID: %s) dengan template: %s", peserta.nama, peserta.id, template_path)

                    # Buat dokumen sementara untuk peserta ini
                    temp_doc = Document(template_path)

                    # Data untuk mengisi placeholder
                    data_dict = {
                        "Nama": peserta.nama or "-",
                        "Jenis_Kelamin": peserta.jenis_kelamin or "-",
                        "NIK": peserta.nik or "-",
                        "Tempat_Lahir": peserta.tempat_lahir or "-",
                        "Tanggal_Lahir": format_tanggal_indonesia(peserta.tanggal_lahir) if peserta.tanggal_lahir else "-",
                        "NISN": peserta.nisn or "-",
                        "Agama_LKP": peserta.agama_lkp or "-",
                        "Handphone": peserta.handphone or "-",
                        "Kewarganegaraan": peserta.kewarganegaraan or "-",
                        "Jenis_Tinggal": peserta.jenis_tinggal or "-",
                        "Tanggal_Masuk": format_tanggal_indonesia(peserta.tanggal_masuk) if peserta.tanggal_masuk else "-",
                        "Email": peserta.email or "-",
                        "Nama_Ortu": peserta.nama_ortu or "-",
                        "NIK_Ortu": peserta.nik_ortu or "-",
                        "Pekerjaan_Ortu": peserta.pekerjaan_ortu or "-",
                        "Pendidikan_Ortu": peserta.pendidikan_ortu or "-",
                        "Penghasilan_Ortu": peserta.penghasilan_ortu or "-",
                        "Handphone_Ortu": peserta.handphone_ortu or "-",
                        "Tempat_Lahir_Ortu": peserta.tempat_lahir_ortu or "-",
                        "Tanggal_Lahir_Ortu": format_tanggal_indonesia(peserta.tanggal_lahir_ortu) if peserta.tanggal_lahir_ortu else "-",
                        "Asal": peserta.asal or "-",
                        "Alamat": peserta.alamat or "-",
                        "RT": peserta.rt or "-",
                        "RW": peserta.rw or "-",
                        "Kecamatan": peserta.kecamatan or "-",
                        "Kelurahan": peserta.kelurahan or "-",
                        "Kab_Kota": peserta.kab_kota or "-",
                        "Propinsi": peserta.propinsi or "-",
                        "Nama_Ibu_kandung": peserta.nama_ibu_kandung or "-",
                        "Nama_Ayah": peserta.nama_ayah or "-",
                        "Agama_Kemdikbud": peserta.agama_kemdikbud or "-",
                        "Penerima_KPS": peserta.penerima_kps or "-",
                        "Layak_PIP": peserta.layak_pip or "-",
                        "Penerima_KIP": peserta.penerima_kip or "-",
                        "Kode_Pos": peserta.kode_pos or "-",
                        "Alat_Transportasi": peserta.alat_transportasi or "-",
                        "Pendidikan_Terakhir": peserta.pendidikan_terakhir or "-",
                        "Nama_Lembaga": peserta.nama_lembaga or "-",
                        "Jabatan": peserta.jabatan or "-",
                        "Alamat_Kantor": peserta.alamat_kantor or "-",
                        "Telp_Kantor": peserta.telp_kantor or "-",
                        "Jadwal": jadwal,
                        "TUK": tuk,
                        "Lokasi_Sertif": lokasi_sertif,
                        "Skema": skema,
                        "Asesor": asesor,
                        "Tanggal_Sertif": tanggal_sertif
                    }
                    logger.debug("[DEBUG] Data untuk placeholder: %s", data_dict)

                    # Ganti placeholder di paragraf
                    for paragraph in temp_doc.paragraphs:
                        for key, value in data_dict.items():
                            placeholder = "{{" + key + "}}"
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, str(value))

                    # Ganti placeholder di tabel
                    for table in temp_doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for key, value in data_dict.items():
                                    placeholder = "{{" + key + "}}"
                                    if placeholder in cell.text:
                                        cell.text = cell.text.replace(placeholder, str(value))

                    # Gabungkan ke dokumen final
                    if first_doc:
                        final_doc = temp_doc
                        first_doc = False
                    else:
                        # Tambahkan page break
                        page_break_paragraph = OxmlElement('w:p')
                        run = OxmlElement('w:r')
                        br = OxmlElement('w:br')
                        br.set(qn('w:type'), 'page')
                        run.append(br)
                        page_break_paragraph.append(run)
                        final_doc.element.body.append(page_break_paragraph)

                        for element in temp_doc.element.body:
                            final_doc.element.body.append(element)

                    # Simpan ke LogHistory2 setelah konversi berhasil untuk peserta ini
                    LogHistory2.objects.create(
                        name=peserta.nama,
                        email=peserta.email or "-",
                        handphone=peserta.handphone or "-",
                        city=peserta.kota or "-",
                        upload_date=timezone.now(),
                        course_name=jadwal,
                        status="Converted",
                        process_time=timezone.now(),
                        file_id=peserta.id,
                        jadwal=jadwal,
                        tuk=tuk,
                        skema=skema,
                        asesor=asesor,
                        lokasi_sertif=lokasi_sertif,
                        template=f"{folder_name}/{os.path.basename(template_path)}"
                    )

                    # Tandai peserta sebagai sudah dikonversi
                    peserta.is_converted = True
                    peserta.save()

                # Simpan dokumen ke BytesIO
                buffer = BytesIO()
                final_doc.save(buffer)
                buffer.seek(0)

                # Simpan ke file sementara dengan nomor urut
                template_name = os.path.basename(template_path).replace('.docx', '')
                filename = f"Sertifikat_{len(peserta_list)}_Peserta_{template_name}_{skema.replace(' ', '_')}.docx"
                temp_file_path = os.path.join(temp_dir, filename)
                with open(temp_file_path, 'wb') as temp_file:
                    temp_file.write(buffer.getvalue())
                buffer.close()

                # Verifikasi file sudah dibuat
                if os.path.exists(temp_file_path):
                    logger.info("[INFO] File sementara berhasil dibuat: %s", temp_file_path)
                else:
                    logger.error("[ERROR] Gagal membuat file sementara: %s", temp_file_path)
                    return JsonResponse({"status": "error", "message": f"Gagal membuat file sementara: {temp_file_path}"}, status=500)

                # Simpan path file untuk pembersihan
                temp_files.append(temp_file_path)

                # Buat URL statis untuk file
                static_url = f"/static/temp/{filename}"
                download_urls.append(static_url)

                # Tambah counter untuk template berikutnya
                counter += 1

            logger.info("[INFO] Berhasil menghasilkan 4 dokumen Word untuk %d peserta dengan skema %s", len(peserta_list), skema)
            return JsonResponse({
                "status": "success",
                "download_urls": download_urls,
                "temp_files": temp_files,
                "message": f"Empat dokumen untuk skema '{skema}' berhasil dihasilkan!"
            })

        except json.JSONDecodeError as e:
            logger.error("[ERROR] Gagal parsing JSON: %s", str(e))
            return JsonResponse({"status": "error", "message": f"Gagal parsing JSON: {str(e)}"}, status=400)
        except Exception as e:
            logger.error("[ERROR] Error konversi: %s\n%s", str(e), traceback.format_exc())
            return JsonResponse({"status": "error", "message": f"Error konversi: {str(e)}"}, status=500)
    else:
        logger.warning("[WARN] Metode tidak diizinkan: %s", request.method)
        return JsonResponse({"status": "error", "message": "Metode tidak diizinkan."}, status=422)

from .models import LogHistory2

@login_required(login_url="login")
def log_history2(request):
    logs = LogHistory2.objects.all()  # Ambil data dari model baru
    for log in logs:
        log.local_time = localtime(log.upload_date).strftime("%d %b %Y %H:%M")
    return render(request, "loghistory2.html", {"logs": logs})

@require_POST
@login_required(login_url="login")
def delete_log2(request, log_id):
    logger.info("[DEBUG] Memproses permintaan delete_log2 untuk log_id: %s", log_id)
    try:
        log = LogHistory2.objects.get(id=log_id)
        log.delete()
        logger.info("[INFO] Berhasil menghapus LogHistory2 dengan ID: %s", log_id)
        return JsonResponse({"status": "success", "message": "Record has been deleted."})
    except LogHistory2.DoesNotExist:
        logger.warning("[WARN] LogHistory2 dengan ID %s tidak ditemukan", log_id)
        return JsonResponse({"status": "error", "message": "Record not found."}, status=404)
    except Exception as e:
        logger.error("[ERROR] Gagal menghapus LogHistory2 dengan ID %s: %s", log_id, str(e))
        return JsonResponse({"status": "error", "message": f"Failed to delete record: {str(e)}"}, status=500)

@login_required(login_url="login")
def download_log2(request, log_id):
    logger.info("[DEBUG] Memproses permintaan download_log2 untuk log_id: %s", log_id)
    if request.method == "POST":
        try:
            # Ambil data log
            log = get_object_or_404(LogHistory2, id=log_id)
            logger.debug("[DEBUG] Data log ditemukan: %s", log.__dict__)

            # Ambil data peserta berdasarkan file_id (jika masih ada) atau gunakan data dari LogHistory2
            peserta = None
            if log.file_id:
                peserta = get_object_or_404(Peserta, id=log.file_id)
                logger.debug("[DEBUG] Peserta ditemukan dari file_id: %s", peserta.__dict__)
            else:
                logger.warning("[WARN] file_id tidak ada, menggunakan data dari LogHistory2")
                peserta = type('Peserta', (), {
                    'nama': log.name,
                    'email': log.email,
                    'handphone': log.handphone,
                    'kota': log.city,
                    'tempat_lahir': None,
                    'tanggal_lahir': None,
                    'jenis_kelamin': None,
                    'alamat': None,
                    'pendidikan_terakhir': None,
                    'nama_lembaga': None,
                    'jabatan': None,
                    'alamat_kantor': None,
                    'telp_kantor': None
                })()

            # Pemetaan skema ke folder
            skema_to_folder = {
                "Associate Data Analyst": "folder1",
                "Instruktur Junior (KKNI Level III)": "folder2",
                "Junior Information Management": "folder3",
                "Pemasangan Jaringan Komputer": "folder4",
                "Pengelolaan Backup dan Restore Data": "folder5",
                "Pengelolaan Data Aplikasi Perkantoran": "folder6",
                "Pengelolaan Keamanan Data Pengguna": "folder7",
                "Pengelolaan Keamanan Jaringan": "folder8",
            }

            skema = log.skema or "Unknown Skema"
            folder_name = skema_to_folder.get(skema, "default")
            logger.info("[INFO] Menggunakan folder template: %s untuk skema: %s", folder_name, skema)

            # Path ke template berdasarkan skema
            template_paths = [
                os.path.join(os.path.dirname(__file__), 'templates', 'docx', folder_name, 'FormPendaftaran.docx'),
                os.path.join(os.path.dirname(__file__), 'templates', 'docx', folder_name, 'AK01.docx'),
                os.path.join(os.path.dirname(__file__), 'templates', 'docx', folder_name, 'APL02.docx'),
                os.path.join(os.path.dirname(__file__), 'templates', 'docx', folder_name, 'APL01.docx')
            ]
            for path in template_paths:
                if not os.path.exists(path):
                    logger.error("[ERROR] Template Word tidak ditemukan di: %s", path)
                    return JsonResponse({"status": "error", "message": f"Template Word {path} tidak ditemukan."}, status=500)

            # Format Tanggal_Sertif
            tanggal_sertif = format_tanggal_indonesia(timezone.now())
            logger.debug("[DEBUG] Tanggal_Sertif: %s", tanggal_sertif)

            # Gunakan data form yang disimpan di LogHistory2
            jadwal = log.jadwal or "No Schedule"
            tuk = log.tuk or "Default TUK"
            asesor = log.asesor or "Default Asesor"
            lokasi_sertif = log.lokasi_sertif or log.city or "Default Location"

            # Buat direktori untuk menyimpan file sementara
            temp_dir = os.path.join(settings.STATICFILES_DIRS[0], 'temp')
            if not os.path.exists(temp_dir):
                os.makedirs(temp_dir)
                logger.info("[INFO] Membuat direktori temp: %s", temp_dir)

            # Bersihkan file lama
            now = time.time()
            for filename in os.listdir(temp_dir):
                file_path = os.path.join(temp_dir, filename)
                if os.path.exists(file_path) and os.stat(file_path).st_mtime < now - 3600:  # 1 jam
                    os.remove(file_path)
                    logger.info("[INFO] Menghapus file lama: %s", file_path)

            # Proses dokumen untuk kedua template
            download_urls = []
            temp_files = []
            counter = 1  # Inisialisasi counter untuk nomor urut

            for template_path in template_paths:
                final_doc = Document(template_path)

                # Data untuk mengisi placeholder
                data_dict = {
                    "Nama": peserta.nama or "-",
                    "Jenis_Kelamin": getattr(peserta, 'jenis_kelamin', '-') or "-",
                    "NIK": getattr(peserta, 'nik', '-') or "-",
                    "Tempat_Lahir": getattr(peserta, 'tempat_lahir', '-') or "-",
                    "Tanggal_Lahir": peserta.tanggal_lahir.strftime("%d %B %Y") if getattr(peserta, 'tanggal_lahir', None) else "-",
                    "NISN": getattr(peserta, 'nisn', '-') or "-",
                    "Agama_LKP": getattr(peserta, 'agama_lkp', '-') or "-",
                    "Handphone": getattr(peserta, 'handphone', '-') or "-",
                    "Kewarganegaraan": getattr(peserta, 'kewarganegaraan', '-') or "-",
                    "Jenis_Tinggal": getattr(peserta, 'jenis_tinggal', '-') or "-",
                    "Tanggal_Masuk": peserta.tanggal_masuk.strftime("%d %B %Y") if getattr(peserta, 'tanggal_masuk', None) else "-",
                    "Email": getattr(peserta, 'email', '-') or "-",
                    "Nama_Ortu": getattr(peserta, 'nama_ortu', '-') or "-",
                    "NIK_Ortu": getattr(peserta, 'nik_ortu', '-') or "-",
                    "Pekerjaan_Ortu": getattr(peserta, 'pekerjaan_ortu', '-') or "-",
                    "Pendidikan_Ortu": getattr(peserta, 'pendidikan_ortu', '-') or "-",
                    "Penghasilan_Ortu": getattr(peserta, 'penghasilan_ortu', '-') or "-",
                    "Handphone_Ortu": getattr(peserta, 'handphone_ortu', '-') or "-",
                    "Tempat_Lahir_Ortu": getattr(peserta, 'tempat_lahir_ortu', '-') or "-",
                    "Tanggal_Lahir_Ortu": peserta.tanggal_lahir_ortu.strftime("%d %B %Y") if getattr(peserta, 'tanggal_lahir_ortu', None) else "-",
                    "Asal": getattr(peserta, 'asal', '-') or "-",
                    "Alamat": getattr(peserta, 'alamat', '-') or "-",
                    "RT": getattr(peserta, 'rt', '-') or "-",
                    "RW": getattr(peserta, 'rw', '-') or "-",
                    "Kecamatan": getattr(peserta, 'kecamatan', '-') or "-",
                    "Kelurahan": getattr(peserta, 'kelurahan', '-') or "-",
                    "Kab_Kota": getattr(peserta, 'kab_kota', '-') or "-",
                    "Propinsi": getattr(peserta, 'propinsi', '-') or "-",
                    "Nama_Ibu_kandung": getattr(peserta, 'nama_ibu_kandung', '-') or "-",
                    "Nama_Ayah": getattr(peserta, 'nama_ayah', '-') or "-",
                    "Agama_Kemdikbud": getattr(peserta, 'agama_kemdikbud', '-') or "-",
                    "Penerima_KPS": getattr(peserta, 'penerima_kps', '-') or "-",
                    "Layak_PIP": getattr(peserta, 'layak_pip', '-') or "-",
                    "Penerima_KIP": getattr(peserta, 'penerima_kip', '-') or "-",
                    "Kode_Pos": getattr(peserta, 'kode_pos', '-') or "-",
                    "Alat_Transportasi": getattr(peserta, 'alat_transportasi', '-') or "-",
                    "Pendidikan_Terakhir": getattr(peserta, 'pendidikan_terakhir', '-') or "-",
                    "Nama_Lembaga": getattr(peserta, 'nama_lembaga', '-') or "-",
                    "Jabatan": getattr(peserta, 'jabatan', '-') or "-",
                    "Alamat_Kantor": getattr(peserta, 'alamat_kantor', '-') or "-",
                    "Telp_Kantor": getattr(peserta, 'telp_kantor', '-') or "-",
                    "Jadwal": jadwal,
                    "TUK": tuk,
                    "Lokasi_Sertif": lokasi_sertif,
                    "Skema": skema,
                    "Asesor": asesor,
                    "Tanggal_Sertif": tanggal_sertif
                }
                logger.debug("[DEBUG] Data untuk placeholder: %s", data_dict)

                # Ganti placeholder di paragraf
                for paragraph in final_doc.paragraphs:
                    for key, value in data_dict.items():
                        placeholder = "{{" + key + "}}"
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, str(value))

                # Ganti placeholder di tabel
                for table in final_doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for key, value in data_dict.items():
                                placeholder = "{{" + key + "}}"
                                if placeholder in cell.text:
                                    cell.text = cell.text.replace(placeholder, str(value))

                # Simpan dokumen ke BytesIO
                buffer = BytesIO()
                final_doc.save(buffer)
                buffer.seek(0)

                # Simpan ke file sementara dengan nomor urut
                template_basename = os.path.basename(template_path).replace('.docx', '')
                filename = f"Sertifikat_{log.name}_{template_basename}_({counter})_{skema.replace(' ', '_')}.docx"
                temp_file_path = os.path.join(temp_dir, filename)
                with open(temp_file_path, 'wb') as temp_file:
                    temp_file.write(buffer.getvalue())
                buffer.close()

                # Simpan path file untuk pembersihan
                temp_files.append(temp_file_path)

                # Buat URL statis untuk file
                static_url = f"/static/temp/{filename}"
                download_urls.append(static_url)

                # Tambah counter untuk template berikutnya
                counter += 1

            logger.info("[INFO] Berhasil menghasilkan 4 dokumen Word untuk log_id: %s dengan skema %s", log_id, skema)
            return JsonResponse({
                "status": "success",
                "download_urls": download_urls,
                "temp_files": temp_files,
                "message": f"Empat dokumen untuk skema '{skema}' berhasil dihasilkan!"
            })

        except Exception as e:
            logger.error("[ERROR] Error download_log2: %s\n%s", str(e), traceback.format_exc())
            return JsonResponse({"status": "error", "message": f"Error mengunduh dokumen: {str(e)}"}, status=500)
    else:
        logger.warning("[WARN] Metode tidak diizinkan: %s", request.method)
        return JsonResponse({"status": "error", "message": "Metode tidak diizinkan."}, status=405)

    # views.py
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse
from django.views.decorators.http import require_POST
from .models import LogHistory2
import logging

logger = logging.getLogger(__name__)


@require_POST
@login_required(login_url="login")
def delete_all_logs2(request):
    logger.info("[DEBUG] Memproses permintaan delete_all_logs2 pada: %s", request.path)
    try:
        deleted_count, _ = LogHistory2.objects.all().delete()
        logger.info("[INFO] Berhasil menghapus %s record LogHistory2", deleted_count)
        return JsonResponse(
            {"status": "success", "message": f"{deleted_count} record telah dihapus."}
        )
    except Exception as e:
        logger.error("[ERROR] Gagal menghapus semua record LogHistory2: %s", str(e))
        return JsonResponse(
            {"status": "error", "message": f"Gagal menghapus record: {str(e)}"},
            status=500,
        )

def format_tanggal_indonesia(tanggal):
    """Mengonversi tanggal ke format DD MMMM YYYY dalam bahasa Indonesia."""
    bulan_indonesia = {
        "January": "Januari",
        "February": "Februari",
        "March": "Maret",
        "April": "April",
        "May": "Mei",
        "June": "Juni",
        "July": "Juli",
        "August": "Agustus",
        "September": "September",
        "October": "Oktober",
        "November": "November",
        "December": "Desember",
    }
    formatted_date = tanggal.strftime("%d %B %Y")
    for eng, indo in bulan_indonesia.items():
        formatted_date = formatted_date.replace(eng, indo)
    # Tambahkan leading zero pada hari jika perlu
    day = formatted_date.split()[0]
    if len(day) == 1:
        formatted_date = f"0{day} {formatted_date.split(' ', 1)[1]}"
    return formatted_date

@require_POST
@login_required(login_url="login")
def cleanup_temp_files(request):
    try:
        data = json.loads(request.body)
        temp_files = data.get("temp_files", [])
        logger.debug("[DEBUG] Menerima permintaan cleanup untuk file: %s", temp_files)

        for temp_file in temp_files:
            if os.path.exists(temp_file):
                os.remove(temp_file)
                logger.info("[INFO] Berhasil menghapus file sementara: %s", temp_file)
            else:
                logger.warning("[WARN] File sementara tidak ditemukan: %s", temp_file)

        return JsonResponse({"status": "success", "message": "File sementara berhasil dihapus."})
    except Exception as e:
        logger.error("[ERROR] Gagal menghapus file sementara: %s", str(e))
        return JsonResponse({"status": "error", "message": f"Gagal menghapus file: {str(e)}"}, status=500)

from .models import ManualPeserta

@login_required(login_url="login")
def input_and_generate(request):
    peserta_list = ManualPeserta.objects.all().order_by('nama')
    return render(request, "input_and_generate.html", {
        'peserta_list': peserta_list
    })

@login_required(login_url="login")
def add_peserta(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            nama = data.get("nama")
            jenis_kelamin = data.get("jenis_kelamin")
            nik = data.get("nik")
            tempat_lahir = data.get("tempat_lahir")
            tanggal_lahir = data.get("tanggal_lahir")
            nisn = data.get("nisn")
            handphone = data.get("handphone")
            email = data.get("email")
            alamat = data.get("alamat")
            kota = data.get("kota")
            kode_pos = data.get("kode_pos")
            pendidikan_terakhir = data.get("pendidikan_terakhir")
            nama_lembaga = data.get("nama_lembaga")
            jabatan = data.get("jabatan")
            alamat_kantor = data.get("alamat_kantor")
            telp_kantor = data.get("telp_kantor")

            if not nama or not jenis_kelamin or not handphone:
                return JsonResponse({"status": "error", "message": "Nama, Jenis Kelamin, dan Handphone wajib diisi."}, status=400)

            peserta = ManualPeserta.objects.create(
                nama=nama,
                jenis_kelamin=jenis_kelamin,
                nik=nik or None,
                tempat_lahir=tempat_lahir or None,
                tanggal_lahir=tanggal_lahir or None,
                nisn=nisn or None,
                handphone=handphone,
                email=email or None,
                alamat=alamat or None,
                kota=kota or None,
                kode_pos=kode_pos or None,
                pendidikan_terakhir=pendidikan_terakhir or None,
                nama_lembaga=nama_lembaga or None,
                jabatan=jabatan or None,
                alamat_kantor=alamat_kantor or None,
                telp_kantor=telp_kantor or None,
                is_converted=False
            )

            return JsonResponse({
                "status": "success",
                "message": f"Peserta {nama} berhasil ditambahkan!",
                "peserta": {
                    "id": peserta.id,
                    "nama": peserta.nama,
                    "email": peserta.email,
                    "handphone": peserta.handphone,
                    "kota": peserta.kota
                }
            })

        except Exception as e:
            logger.error("[ERROR] Gagal menambah peserta: %s", str(e))
            return JsonResponse({"status": "error", "message": f"Terjadi kesalahan: {str(e)}"}, status=500)

    return JsonResponse({"status": "error", "message": "Metode tidak diizinkan."}, status=405)

@login_required(login_url="login")
def input_and_generate_delete_peserta(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            peserta_ids = data.get("peserta_ids", [])
            if not peserta_ids:
                return JsonResponse({"status": "error", "message": "Tidak ada peserta yang dipilih."}, status=400)

            deleted_count = ManualPeserta.objects.filter(id__in=peserta_ids).delete()[0]
            return JsonResponse({
                "status": "success",
                "message": f"{deleted_count} peserta berhasil dihapus."
            })

        except Exception as e:
            logger.error("[ERROR] Gagal menghapus peserta: %s", str(e))
            return JsonResponse({"status": "error", "message": f"Terjadi kesalahan: {str(e)}"}, status=500)

    return JsonResponse({"status": "error", "message": "Metode tidak diizinkan."}, status=405)

@login_required(login_url="login")
def input_and_generate_convert_document(request):
    logger.info("[DEBUG] Memproses permintaan input_and_generate_convert_document pada: %s", request.path)
    if request.method == "POST":
        try:
            # Langkah 1: Parse request body
            logger.debug("[DEBUG] Request body: %s", request.body)
            data = json.loads(request.body)
            peserta_ids = data.get("peserta_ids", [])
            jadwal = data.get("jadwal")
            tuk = data.get("tuk")
            skema = data.get("skema")
            asesor = data.get("asesor")
            lokasi_sertif = data.get("lokasi_sertif")
            logger.debug("[DEBUG] Data diterima: peserta_ids=%s, jadwal=%s, tuk=%s, skema=%s, asesor=%s, lokasi_sertif=%s",
                         peserta_ids, jadwal, tuk, skema, asesor, lokasi_sertif)

            # Langkah 2: Validasi input
            if not peserta_ids:
                logger.warning("[ERROR] Tidak ada peserta dipilih untuk konversi")
                return JsonResponse({"status": "error", "message": "Pilih setidaknya satu peserta."}, status=400)

            if not all([jadwal, tuk, skema, asesor, lokasi_sertif]):
                logger.warning("[ERROR] Field form tidak lengkap")
                return JsonResponse({"status": "error", "message": "Semua field form harus diisi."}, status=400)

            # Langkah 3: Mengambil data dari tabel ManualPeserta
            logger.debug("[INFO] Mengambil peserta dengan ID: %s", peserta_ids)
            peserta_list = ManualPeserta.objects.filter(id__in=peserta_ids)
            if not peserta_list.exists():
                logger.error("[ERROR] Peserta tidak ditemukan: %s", peserta_ids)
                return JsonResponse({"status": "error", "message": "Peserta tidak ditemukan."}, status=404)

            # Langkah 4: Membuat format tanggal_sertif
            tanggal_sertif = format_tanggal_indonesia(timezone.now())
            logger.debug("[DEBUG] Tanggal_Sertif: %s", tanggal_sertif)

            # Langkah 5: Menentukan folder berdasarkan skema
            skema_to_folder = {
                "Associate Data Analyst": "folder1",
                "Instruktur Junior (KKNI Level III)": "folder2",
                "Junior Information Management": "folder3",
                "Pemasangan Jaringan Komputer": "folder4",
                "Pengelolaan Backup dan Restore Data": "folder5",
                "Pengelolaan Data Aplikasi Perkantoran": "folder6",
                "Pengelolaan Keamanan Data Pengguna": "folder7",
                "Pengelolaan Keamanan Jaringan": "folder8",
            }

            if skema not in skema_to_folder:
                logger.error("[ERROR] Skema tidak valid: %s", skema)
                return JsonResponse({"status": "error", "message": f"Skema '{skema}' tidak valid."}, status=400)

            folder_name = skema_to_folder[skema]
            logger.info("[INFO] Menggunakan folder template: %s untuk skema: %s", folder_name, skema)

            # Langkah 6: Menentukan path ke template Word
            template_paths = [
                os.path.join(os.path.dirname(__file__), 'templates', 'docx', folder_name, 'FormPendaftaran.docx'),
                os.path.join(os.path.dirname(__file__), 'templates', 'docx', folder_name, 'AK01.docx'),
                os.path.join(os.path.dirname(__file__), 'templates', 'docx', folder_name, 'APL02.docx'),
                os.path.join(os.path.dirname(__file__), 'templates', 'docx', folder_name, 'APL01.docx')
            ]
            for template_path in template_paths:
                if not os.path.exists(template_path):
                    logger.error("[ERROR] Template Word tidak ada di: %s", template_path)
                    return JsonResponse({"status": "error", "message": f"Template Word {template_path}: tidak ditemukan."}, status=500)

            # Membuat direktori sementara di STATICFILES_DIRS
            temp_dir = os.path.join(settings.STATICFILES_DIRS[0], 'temp')
            if not os.path.exists(temp_dir):
                os.makedirs(temp_dir)
                logger.info("[INFO] Direktori temp dibuat: %s", temp_dir)

            # Membersihkan file lama
            now = time.time()
            for filename in os.listdir(temp_dir):
                file_path = os.path.join(temp_dir, filename)
                if os.path.isfile(file_path) and os.stat(file_path).st_mtime < now - 3600:  # 1 jam
                    os.remove(file_path)
                    logger.info("[INFO] File lama dihapus: %s", file_path)

            # Langkah 7: Memproses peserta dan membuat dokumen
            download_urls = []
            temp_files = []
            counter = 1  # Inisialisasi counter untuk nomor urut

            for template_path in template_paths:
                final_doc = Document()
                first_doc = True

                for index, peserta in enumerate(peserta_list):
                    logger.debug("[INFO] Memproses peserta: %s (ID: %d) dengan template %s", peserta.nama, peserta.id, template_path)

                    # Mengambil template dokumen
                    temp_doc = Document(template_path)

                    # Data untuk placeholder
                    data_dict = {
                        "Nama": peserta.nama or "-",
                        "Jenis_Kelamin": peserta.jenis_kelamin or "-",
                        "NIK": peserta.nik or "-",
                        "Tempat_Lahir": peserta.tempat_lahir or "-",
                        "Tanggal_Lahir": format_tanggal_indonesia(peserta.tanggal_lahir) if peserta.tanggal_lahir else "-",
                        "NISN": peserta.nisn or "-",
                        "Handphone": peserta.handphone or "-",
                        "Email": peserta.email or "-",
                        "Alamat": peserta.alamat or "-",
                        "Kota": peserta.kota or "-",
                        "KodePos": peserta.kode_pos or "-",
                        "Pendidikan_Terakhir": peserta.pendidikan_terakhir or "-",
                        "Nama_Lembaga": peserta.nama_lembaga or "-",
                        "Jabatan": peserta.jabatan or "-",
                        "Alamat_Kantor": peserta.alamat_kantor or "-",
                        "Telp_Kantor": peserta.telp_kantor or "-",
                        "Agama_LKP": getattr(peserta, "agama_lkp", "-"),
                        "Kewarganegaraan": getattr(peserta, "kewarganegaraan", "-"),
                        "Jenis_Tinggal": getattr(peserta, "jenis_tinggal", "-"),
                        "Tanggal_Masuk": "-",
                        "Nama_Ortu": getattr(peserta, "nama_ortu", "-"),
                        "NIK_Ortu": getattr(peserta, "nik_ortu", "-"),
                        "Pekerjaan_Ortu": getattr(peserta, "pekerjaan_ortu", "-"),
                        "Pendidikan_Ortu": getattr(peserta, "pendidikan_ortu", "-"),
                        "Penghasilan_Ortu": getattr(peserta, "penghasilan_ortu", "-"),
                        "Handphone_Ortu": getattr(peserta, "handphone_ortu", "-"),
                        "Tempat_Lahir_Ortu": getattr(peserta, "tempat_lahir_ortu", "-"),
                        "Tanggal_Lahir_Ortu": format_tanggal_indonesia(peserta.tanggal_lahir_ortu) if getattr(peserta, "tanggal_lahir_ortu", None) else "-",
                        "Asal": getattr(peserta, "asal", "-"),
                        "RT": getattr(peserta, "rt", "-"),
                        "RW": getattr(peserta, "rw", "-"),
                        "Kecamatan": getattr(peserta, "kecamatan", "-"),
                        "Kelurahan": getattr(peserta, "kelurahan", "-"),
                        "Kab_Kota": getattr(peserta, "kab_kota", "-"),
                        "Propinsi": getattr(peserta, "propinsi", "-"),
                        "Nama_Ibu_Kandung": getattr(peserta, "nama_ibu_kandung", "-"),
                        "Nama_Ayah": getattr(peserta, "nama_ayah", "-"),
                        "Agama_Kemdikbud": getattr(peserta, "agama_kemdikbud", "-"),
                        "Penerima_KPS": getattr(peserta, "penerima_kps", "-"),
                        "Layak_PIP": getattr(peserta, "layak_pip", "-"),
                        "Penerima_KIP": getattr(peserta, "penerima_kip", "-"),
                        "Alat_Transportasi": getattr(peserta, "alat_transportasi", "-"),
                        "Jadwal": jadwal,
                        "TUK": tuk,
                        "Lokasi_Sertif": lokasi_sertif,
                        "Skema": skema,
                        "Asesor": asesor,
                        "Tanggal_Sertif": tanggal_sertif
                    }
                    # Ganti placeholder di paragraf
                    for paragraph in temp_doc.paragraphs:
                        for key, value in data_dict.items():
                            placeholder = "{{" + key + "}}"
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, str(value))

                    # Ganti placeholder dalam tabel
                    for table in temp_doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for key, value in data_dict.items():
                                    placeholder = f"{{{{{key}}}}}"
                                    if placeholder in cell.text:
                                        cell.text = cell.text.replace(placeholder, str(value))

                    # Gabungkan dokumen
                    if first_doc:
                        final_doc = temp_doc
                        first_doc = False
                    else:
                        page_break_paragraph = OxmlElement('w:p')
                        run = OxmlElement('w:r')
                        br = OxmlElement('w:br')
                        br.set(qn('w:type'), 'page')
                        run.append(br)
                        page_break_paragraph.append(run)
                        final_doc.element.body.append(page_break_paragraph)

                        for element in temp_doc.element.body:
                            final_doc.element.body.append(element)

                    # Simpan log ke tabel LogHistory2
                    LogHistory2.objects.create(
                        name=peserta.nama,
                        email=peserta.email or "-",
                        handphone=peserta.handphone or "-",
                        city=peserta.kota or "-",
                        upload_date=timezone.now(),
                        status="Converted",
                        process_time=timezone.now(),
                        file_id=peserta.id,
                        jadwal=jadwal,
                        tuk=tuk,
                        skema=skema,
                        asesor=asesor,
                        lokasi_sertif=lokasi_sertif,
                        template=f"{folder_name}/{os.path.basename(template_path)}"
                    )

                    # Tandai peserta sebagai sudah dikonversi
                    peserta.is_converted = True
                    peserta.save()

                # Simpan dokumen ke BytesIO
                buffer = BytesIO()
                final_doc.save(buffer)
                buffer.seek(0)

                # Simpan file sementara
                template_name = os.path.basename(template_path).replace('.docx', '')
                filename = f"Sertifikasi_{len(peserta_list)}_Peserta_{template_name}_({counter})_{skema.replace(' ', '_')}.docx"
                temp_file_path = os.path.join(temp_dir, filename)
                with open(temp_file_path, 'wb') as temp_file:
                    temp_file.write(buffer.getvalue())
                buffer.close()

                # Verifikasi file
                if os.path.exists(temp_file_path):
                    logger.info("[INFO] File sementara dibuat: %s", temp_file_path)
                else:
                    logger.error("[ERROR] Gagal membuat file: %s", temp_file_path)
                    return JsonResponse({"status": "error", "message": f"Gagal membuat file: {temp_file_path}"}, status=500)

                temp_files.append(temp_file_path)
                static_url = f"/static/temp/{filename}"
                download_urls.append(static_url)

                # Tambah counter untuk template berikutnya
                counter += 1

            logger.info("[INFO] Berhasil menghasilkan 4 dokumen untuk %d peserta dengan skema %s", len(peserta_list), skema)
            return JsonResponse({
                "status": "success",
                "download_urls": download_urls,
                "temp_files": temp_files,
                "message": f"Empat dokumen untuk skema '{skema}' berhasil dihasilkan!"
            })

        except json.JSONDecodeError as e:
            logger.error("[ERROR] Gagal parsing JSON: %s", str(e))
            return JsonResponse({"status": "error", "message": f"Gagal parsing JSON: {str(e)}"}, status=400)
        except Exception as e:
            logger.error("[ERROR] Error konversi: %s\n%s", str(e), traceback.format_exc())
            return JsonResponse({"status": "error", "message": str(e)}, status=500)
    else:
        logger.warning("[WARN] Metode tidak diizinkan: %s", request.method)
        return JsonResponse({"status": "error", "message": "Metode tidak diizinkan."}, status=405)

@require_POST
@login_required(login_url="login")
def input_and_generate_cleanup_temp_files(request):
    logger.info("[DEBUG] Memproses permintaan input_and_generate_cleanup_temp_files pada: %s", request.path)
    try:
        data = json.loads(request.body)
        temp_files = data.get("temp_files", [])
        logger.debug("[DEBUG] Menerima permintaan cleanup untuk file: %s", temp_files)

        temp_dir = os.path.join(settings.STATICFILES_DIRS[0], 'temp')
        for temp_file in temp_files:
            # Pastikan file berada di direktori temp untuk keamanan
            temp_file_path = os.path.join(temp_dir, os.path.basename(temp_file))
            if os.path.exists(temp_file_path):
                os.remove(temp_file_path)
                logger.info("[INFO] Berhasil menghapus file sementara: %s", temp_file_path)
            else:
                logger.warning("[WARN] File sementara tidak ditemukan: %s", temp_file_path)

        return JsonResponse({"status": "success", "message": "File sementara berhasil dihapus."})
    except json.JSONDecodeError as e:
        logger.error("[ERROR] Gagal parsing JSON: %s", str(e))
        return JsonResponse({"status": "error", "message": f"Data request tidak valid: {str(e)}"}, status=400)
    except Exception as e:
        logger.error("[ERROR] Gagal menghapus file sementara: %s", str(e))
        return JsonResponse({"status": "error", "message": f"Gagal menghapus file: {str(e)}"}, status=500)

@login_required(login_url="login")
def input_data_status(request):
    files = UploadedFile.objects.all()
    return render(request, 'input_data_status.html', {'files': files})